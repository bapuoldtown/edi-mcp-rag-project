"""
Word document parser (.docx)
"""

import logging
from typing import List

from .base_parser import BaseParser
from .models import ParsedDocument, DocumentType, Table

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parse Word documents (.docx)"""
    
    SUPPORTED_EXTENSIONS = ['docx', 'doc']
    
    def can_parse(self, file_path: str) -> bool:
        """Check if file is Word document"""
        return self.get_file_extension(file_path) in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> ParsedDocument:
        """Parse Word document"""
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid file: {file_path}")
        
        ext = self.get_file_extension(file_path)
        
        if ext == 'doc':
            logger.warning(".doc format not fully supported, results may be limited")
            return self._parse_as_text(file_path)
        
        logger.info(f"Parsing Word document: {file_path}")
        
        try:
            return self._parse_docx(file_path)
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            raise
    
    def _parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse .docx file"""
        from docx import Document
        
        doc = Document(file_path)
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        combined_text = "\n\n".join(paragraphs)
        
        # Extract tables if enabled
        tables = []
        if self.config.extract_tables and doc.tables:
            for table in doc.tables:
                parsed_table = self._parse_word_table(table)
                if parsed_table.rows:
                    tables.append(parsed_table)
        
        # Metadata
        metadata = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables)
        }
        
        # Try to get core properties
        try:
            core_props = doc.core_properties
            metadata.update({
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else ''
            })
        except:
            pass
        
        return self.create_parsed_document(
            file_path=file_path,
            document_type=DocumentType.WORD,
            text=combined_text,
            tables=tables,
            metadata=metadata
        )
    
    def _parse_word_table(self, word_table) -> Table:
        """Parse a Word table"""
        rows_data = []
        
        for row in word_table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip()
                row_data.append(text)
            rows_data.append(row_data)
        
        if not rows_data:
            return Table()
        
        # First row as headers
        headers = rows_data[0]
        rows = rows_data[1:]
        
        return Table(
            headers=headers,
            rows=rows,
            num_rows=len(rows),
            num_columns=len(headers)
        )
    
    def _parse_as_text(self, file_path: str) -> ParsedDocument:
        """Fallback for .doc files"""
        return self.create_parsed_document(
            file_path=file_path,
            document_type=DocumentType.WORD,
            text="Note: .doc format requires conversion to .docx for full parsing",
            metadata={'warning': 'Limited parsing for .doc format'}
        )
